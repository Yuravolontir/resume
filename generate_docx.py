from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

# Colors
ACCENT = RGBColor(0x94, 0x33, 0x19)
TEXT = RGBColor(0x1b, 0x1c, 0x1c)
TEXT_2 = RGBColor(0x57, 0x42, 0x3d)
MONO = 'Consolas'
BODY = 'Calibri'
HEADING = 'Georgia'

style = doc.styles['Normal']
font = style.font
font.name = BODY
font.size = Pt(11)
font.color.rgb = TEXT

def add_horizontal_line(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'DEC0B9')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_section_header(doc, number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run_num = p.add_run(number + '  ')
    run_num.font.name = MONO
    run_num.font.size = Pt(9)
    run_num.font.color.rgb = ACCENT
    run_title = p.add_run(title)
    run_title.font.name = HEADING
    run_title.font.size = Pt(16)
    run_title.font.color.rgb = TEXT
    run_title.bold = True
    add_horizontal_line(doc)

def add_skill_chip(paragraph, text):
    run = paragraph.add_run(text + '   ')
    run.font.size = Pt(9.5)
    run.font.name = BODY
    run.font.color.rgb = TEXT_2

# ═══════════════════════════════════════
# HERO / HEADER
# ═══════════════════════════════════════
# Name
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Yura Volontir')
run.font.name = HEADING
run.font.size = Pt(26)
run.font.color.rgb = TEXT

# Label
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run('JUNIOR FULL-STACK DEVELOPER')
run.font.name = MONO
run.font.size = Pt(9)
run.font.color.rgb = ACCENT

# Title
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
run = p.add_run('Full-stack developer building applications with React, Node.js, C#, and AI/ML')
run.font.name = BODY
run.font.size = Pt(11)
run.font.italic = True
run.font.color.rgb = TEXT_2

# Summary
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
run = p.add_run('Full-stack developer with experience in web development, desktop applications, database design, neural networks, and RAG systems. Focused on building clean, functional software.')
run.font.name = BODY
run.font.size = Pt(10.5)
run.font.color.rgb = TEXT_2

# Tags
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
for tag in ['React', 'Node.js', 'C# / .NET', 'SQL Server', 'Android', 'Python', 'AI / ML']:
    run = p.add_run(tag + '    ')
    run.font.size = Pt(9)
    run.font.name = BODY
    run.font.color.rgb = TEXT_2

# Contact
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
contacts = [
    'yuravolontir@gmail.com',
    '053-4352832',
    'Netanya, Israel',
    'linkedin.com/in/yura-volontir',
    'github.com/Yuravolontir',
]
run = p.add_run('  |  '.join(contacts))
run.font.size = Pt(9)
run.font.color.rgb = TEXT_2

add_horizontal_line(doc)

# ═══════════════════════════════════════
# EDUCATION
# ═══════════════════════════════════════
add_section_header(doc, '01', 'Education')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Ruppin College (Machon Atid)')
run.font.name = HEADING
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = TEXT

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('Software Engineering — Practical Engineer')
run.font.size = Pt(11)
run.font.color.rgb = ACCENT
run.bold = True

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('2024 – 2026')
run.font.name = MONO
run.font.size = Pt(9)
run.font.color.rgb = ACCENT

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Full-stack web development (React, Node.js, Express), C# / .NET ecosystem, Android development, database design & administration (SQL Server, MongoDB), AI/ML (neural networks, NLP, RAG with LangChain), OOP, data structures & algorithms, systems analysis, and cybersecurity.')
run.font.size = Pt(10)
run.font.color.rgb = TEXT_2

# ═══════════════════════════════════════
# SKILLS
# ═══════════════════════════════════════
add_section_header(doc, '02', 'Technical Skills')

skills = [
    ('Languages', 'C#, JavaScript ES6+, TypeScript, Python, Java, SQL (T-SQL), HTML5, CSS3'),
    ('Frontend', 'React 19, JSX, Craft.js, Tailwind CSS, Material UI, WinForms, WPF / XAML'),
    ('Backend', 'Node.js, Express.js, REST API, .NET 8, CRUD'),
    ('Databases', 'SQL Server (Stored Procedures, Triggers), MongoDB, ChromaDB, ER Diagrams'),
    ('AI & ML', 'Neural Networks, NLP, LSTM, RAG, LangChain, Ollama, brain.js'),
    ('Tools', 'Git / GitHub, Visual Studio, VS Code, Android Studio, Vite, Netlify'),
]

for name, chips in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(3)
    run = p.add_run(name + ':  ')
    run.font.name = HEADING
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = TEXT
    run = p.add_run(chips)
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_2

# ═══════════════════════════════════════
# PROJECTS
# ═══════════════════════════════════════
add_section_header(doc, '03', 'Projects')

projects = [
    {
        'name': 'Visual Website Builder',
        'link': 'github.com/Yuravolontir/WebBuildReactExam',
        'demo': 'dragcanvasapp.netlify.app',
        'desc': 'Full-stack drag-and-drop website builder with a visual editor powered by Craft.js. Component library, real-time editing, Express.js backend, MSSQL persistence.',
        'tech': ['React 19', 'Express.js', 'MSSQL', 'Craft.js'],
    },
    {
        'name': 'MyLibrary',
        'link': 'github.com/Yuravolontir/MyLibrary',
        'desc': 'Desktop library management system with user authentication (SHA256), role-based access control (admin/user), book CRUD operations, reading time tracker, and top readers leaderboard. All database operations via SQL Server stored procedures.',
        'tech': ['C#', 'WinForms', 'SQL Server', 'Stored Procedures'],
    },
    {
        'name': 'AI Design Assistant',
        'link': 'github.com/Yuravolontir/Ai-Design-Asistant',
        'desc': 'AI-powered design tool that recommends color palettes and slogans based on project type and mood. Uses neural networks (brain.js) for color classification and LSTM for slogan generation, trained on custom datasets.',
        'tech': ['JavaScript', 'brain.js', 'Neural Networks', 'LSTM'],
    },
]

for proj in projects:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(proj['name'])
    run.font.name = HEADING
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = TEXT

    # Link
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(proj['link'])
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT
    if 'demo' in proj:
        run2 = p.add_run('  |  Live: ' + proj['demo'])
        run2.font.size = Pt(9)
        run2.font.color.rgb = ACCENT

    # Description
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(proj['desc'])
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT_2

    # Tech tags
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for t in proj['tech']:
        run = p.add_run(t + '   ')
        run.font.name = MONO
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x83, 0x26, 0x0d)
        run.bold = True

# ═══════════════════════════════════════
# CERTIFICATIONS & LANGUAGES
# ═══════════════════════════════════════
add_section_header(doc, '04', 'Certifications & Languages')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run('CCNA: Introduction to Networks')
run.font.size = Pt(11)
run.font.bold = True
p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(10)
run = p2.add_run('Cisco Networking Academy  —  Certified')
run.font.size = Pt(10)
run.font.color.rgb = TEXT_2

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Languages:  ')
run.font.name = HEADING
run.font.size = Pt(12)
run.font.bold = True
run = p.add_run('Russian (Native)  |  Hebrew (Professional Working)  |  English (Professional Working)')
run.font.size = Pt(10.5)
run.font.color.rgb = TEXT_2

# Save
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Yura_Volontir_Resume.docx')
doc.save(output_path)
print('Saved:', output_path)
